import os
import time
from utils import get_project_root

OUTPUT_DIR = os.path.abspath(os.path.join(get_project_root(), "outputs", "files"))


def ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def create_docx(title: str, text_content: str, filename: str = None) -> str:
    """
    Generates a Word DOCX file containing the provided Title and Text Content.
    Returns the absolute path to the generated file.
    """
    # Import python-docx (installed as 'python-docx', imported as 'docx')
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    ensure_dir(OUTPUT_DIR)

    if not filename:
        timestamp = int(time.time())
        clean_title = "".join([c for c in title if c.isalnum() or c == ' ']).strip().replace(' ', '_')
        filename = f"{clean_title}_{timestamp}.docx"

    if not filename.endswith('.docx'):
        filename += '.docx'

    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    doc.add_heading(title, 0)

    paragraphs = text_content.split('\n')
    for p in paragraphs:
        if p.strip():
            doc.add_paragraph(p.strip())

    doc.save(filepath)
    return os.path.abspath(filepath)


if __name__ == "__main__":
    res = create_docx("Jarvis Diagnostic", "Protocol initialized.\nMemory banks: Stable.")
    print(f"Generated DOCX: {res}")
